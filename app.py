import streamlit as st
import json
import copy
from docx import Document
from io import BytesIO
from pathlib import Path

def replace_text_in_cell(cell, old_text, new_text):
    """
    Helper function to replace text inside a table cell.
    We iterate through paragraphs to preserve the cell structure.
    """
    for paragraph in cell.paragraphs:
        if old_text in paragraph.text:
            # Replace the text within the paragraph
            paragraph.text = paragraph.text.replace(old_text, new_text)

def process_template(template_file, json_data):
    """
    Takes the uploaded Word document and the JSON data, and builds
    the populated document by copying tables and rows.
    """
    doc = Document(template_file)

    # 1. Locate the template table containing the $sub_heading$ tag
    template_table = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "$sub_heading$" in cell.text:
                    template_table = table
                    break
            if template_table:
                break
        if template_table:
            break

    if not template_table:
        st.error("Could not find a table containing '$sub_heading$' in the uploaded template.")
        return None

    # 2. Iterate through the JSON sections to create a new table for each
    sections = json_data.get("revision_sections", [])

    for section_idx, section in enumerate(sections):
        # Add a paragraph with text before adding a new table (except for the first table)
        if section_idx > 0:
            doc.add_paragraph("BLANK LINE")
        
        # Deep copy the entire template table XML and append it to the document body
        new_tbl_xml = copy.deepcopy(template_table._tbl)
        doc.element.body.append(new_tbl_xml)

        # The newly appended table is now the last one in doc.tables
        new_table = doc.tables[-1]

        # Replace the $sub_heading$ tag in the new table
        for row in new_table.rows:
            for cell in row.cells:
                replace_text_in_cell(cell, "$sub_heading$", section.get("sub_heading", ""))

        # 3. Locate the row containing the $prompt$ tag to act as our row template
        prompt_row_idx = -1
        for i, row in enumerate(new_table.rows):
            for cell in row.cells:
                if "$prompt$" in cell.text:
                    prompt_row_idx = i
                    break
            if prompt_row_idx != -1:
                break

        if prompt_row_idx != -1:
            template_row = new_table.rows[prompt_row_idx]
            
            # Save the pristine XML of the row containing the tags BEFORE we overwrite it
            original_row_xml = copy.deepcopy(template_row._tr)
            
            questions = section.get("questions", [])

            for q_idx, q in enumerate(questions):
                prompt_text = q.get("prompt", "")
                # Join the answer space array into a single string with line breaks
                answer_text = "\n".join(q.get("answer_space", []))

                if q_idx == 0:
                    # Modify the existing row for the first question
                    for cell in template_row.cells:
                        replace_text_in_cell(cell, "$prompt$", prompt_text)
                        replace_text_in_cell(cell, "$answer_space$", answer_text)
                else:
                    # Duplicate the row XML for subsequent questions
                    new_row_xml = copy.deepcopy(original_row_xml)
                    new_table._tbl.append(new_row_xml)
                    new_row = new_table.rows[-1]

                    # Replace tags in the newly duplicated row
                    for cell in new_row.cells:
                        replace_text_in_cell(cell, "$prompt$", prompt_text)
                        replace_text_in_cell(cell, "$answer_space$", answer_text)

    # 4. Remove the original template table from the document to clean it up
    template_table._element.getparent().remove(template_table._element)

    # 5. Save to a BytesIO object so Streamlit can offer it as a download
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output

def get_available_templates():
    """
    Get list of available template files (.docx) in the repository root.
    """
    current_dir = Path(".")
    docx_files = sorted([f.name for f in current_dir.glob("*.docx")])
    return docx_files

# --- Streamlit UI ---
st.set_page_config(page_title="Revision Template Generator", layout="centered")

st.title("📄 Revision Template Generator")
st.markdown("Upload your `.docx` template and paste your JSON data below to generate populated revision tables.")

# Template Selection
col1, col2 = st.columns(2)

with col1:
    use_repository_template = st.checkbox("Use template from repository", value=False)

template_file = None

if use_repository_template:
    available_templates = get_available_templates()
    
    if available_templates:
        selected_template = st.selectbox(
            "Select a template from the repository",
            available_templates,
            help="Available templates in the repository root"
        )
        template_path = Path(selected_template)
        try:
            with open(template_path, "rb") as f:
                template_file = BytesIO(f.read())
                st.success(f"✓ Using: {selected_template}")
        except Exception as e:
            st.error(f"Error loading template: {e}")
    else:
        st.warning("No .docx template files found in the repository root. Please upload a template instead.")
        use_repository_template = False

if not use_repository_template:
    # File Uploader for Template
    uploaded_template = st.file_uploader("1. Upload your Template Recall.docx", type=["docx"])
    if uploaded_template:
        template_file = uploaded_template

# Text Area for JSON
json_input = st.text_area("2. Paste your JSON Data", height=300, placeholder='{"revision_sections": [...]}')

if st.button("Generate Document", type="primary"):
    if not template_file:
        st.warning("Please either select a repository template or upload a Word document template.")
    elif not json_input.strip():
        st.warning("Please paste your JSON data.")
    else:
        try:
            # Parse JSON
            parsed_json = json.loads(json_input)
            
            with st.spinner("Generating document..."):
                # Process the document
                output_docx = process_template(template_file, parsed_json)
                
                if output_docx:
                    st.success("Document generated successfully!")
                    
                    # Download Button
                    st.download_button(
                        label="⬇️ Download Populated Word Document",
                        data=output_docx,
                        file_name="Populated_Revision_Document.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        except json.JSONDecodeError:
            st.error("Invalid JSON format. Please check your JSON data and try again.")
        except Exception as e:
            st.error(f"An error occurred: {e}")
